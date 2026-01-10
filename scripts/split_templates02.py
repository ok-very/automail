import os
import re
from docx import Document

def sanitize_filename(name: str) -> str:
    """Replace invalid filename characters"""
    return re.sub(r'[<>:"/\\|?*\n\r]+', '_', name).strip()

def copy_paragraph(src_para, dst_doc):
    """Copy a paragraph with formatting into a new docx document"""
    style_name = src_para.style.name if src_para.style else ""
    # Handle bullet/numbered list
    if "List Bullet" in style_name:
        dst_para = dst_doc.add_paragraph(style='ListBullet')
    elif "List Number" in style_name:
        dst_para = dst_doc.add_paragraph(style='ListNumber')
    else:
        dst_para = dst_doc.add_paragraph()

    # Copy runs (bold, italic, underline, font, size)
    for run in src_para.runs:
        dst_run = dst_para.add_run(run.text)
        dst_run.bold = run.bold
        dst_run.italic = run.italic
        dst_run.underline = run.underline
        dst_run.font.name = run.font.name
        dst_run.font.size = run.font.size

    return dst_para

def write_docx_section(heading, paras, output_dir, seq_number):
    """Write a new DOCX file for a single heading section with numbered filename"""
    doc = Document()
    doc.add_heading(heading, level=3)  # H3 heading now
    for para in paras:
        copy_paragraph(para, doc)
    # Pad sequence number to 2 digits
    seq_str = str(seq_number).zfill(2)
    fname = f"{seq_str}_{sanitize_filename(heading)}.docx"
    out_path = os.path.join(output_dir, fname)
    doc.save(out_path)
    print(f"Wrote: {out_path}")
    return fname

def process_docx_file(input_path, output_dir):
    """Split a single DOCX file into sections based on H3 headings"""
    doc = Document(input_path)
    current_heading = None
    collected_paras = []
    heading_list = []
    seq_number = 1

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ''
        if style_name.startswith("Heading ") and style_name.endswith("3"):  # H3 detection
            # Save previous section
            if current_heading and collected_paras:
                write_docx_section(current_heading, collected_paras, output_dir, seq_number)
                heading_list.append(current_heading)
                seq_number += 1
            current_heading = para.text.strip()
            collected_paras = []
        else:
            if current_heading:
                collected_paras.append(para)

    # Save the last section
    if current_heading and collected_paras:
        write_docx_section(current_heading, collected_paras, output_dir, seq_number)
        heading_list.append(current_heading)

    # Write all headings to a .txt file
    headings_txt_path = os.path.join(output_dir, "headings_list.txt")
    with open(headings_txt_path, "w", encoding="utf-8") as f:
        for i, h in enumerate(heading_list, start=1):
            f.write(f"{i}. {h}\n")
    print(f"All headings written to {headings_txt_path}")

if __name__ == "__main__":
    # === Specify input file and output folder ===
    input_file = "docs/email-templates02.docx"  # <-- change this to your file
    output_dir = "output_docx2"
    os.makedirs(output_dir, exist_ok=True)

    process_docx_file(input_file, output_dir)
    print("Done!")
